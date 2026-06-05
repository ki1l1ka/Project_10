import docx
import os
import io
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
from logic import NuclearDataVisualizer
class Word:
    def __init__(self, filename):
        self.FileName = filename
        self.Extension = 'docx'
        self.FullName = f'{self.FileName}.{self.Extension}'

        if os.path.exists(self.FullName):
            self.doc = docx.Document(self.FullName)
        else:
            self.doc = docx.Document()
            self.doc.save(self.FullName)

    def Save(self):
        self.doc.save(self.FullName)

    def ReadFileText(self):
        text = [par.text for par in self.doc.paragraphs]
        return '\n'.join(text)

    def ReadFileTable(self, table_number):
        if table_number >= len(self.doc.tables):
            return "Таблица не найдена"
        table = self.doc.tables[table_number]
        return [[cell.text for cell in row.cells] for row in table.rows]

    def AddHeader(self, header="Заголовок", level=1):
        self.doc.add_heading(header, level=level)
        self.Save()

    def AddText(self, text='Текст.'):
        self.doc.add_paragraph(text)
        self.Save()

    def AddTable(self, table_obj):
        rows = len(table_obj.matrix)
        cols = len(table_obj.matrix[0])
        word_table = self.doc.add_table(rows=rows, cols=cols)


def ToExcel(filename, table, sheetname=False):
    df = pd.DataFrame(table)
    if not sheetname:
        df.to_excel(f'excel/{filename}.xlsx', index=False)
    else:
        with pd.ExcelWriter(f'excel/{filename}.xlsx') as writer:
            df.to_excel(writer, sheet_name=sheetname)
    return None


class ReportDocument:
    def __init__(self, title="Аналитический отчет по обращению с ОЯТ и РАО"):
        self.doc = Document()
        self.doc.add_heading(title, 0)

        # Папка для сохранения по умолчанию
        self.folder = "отчёты"
        if not os.path.exists(self.folder):
            os.makedirs(self.folder)

    def AddHeading(self, text, level=1):
        self.doc.add_heading(text, level=level)

    def AddParagraph(self, text):
        if text:
            self.doc.add_paragraph(text)

    def AddPageBreak(self):
        self.doc.add_page_break()

    def AddTable(self, table_obj, plant_name=None):
        if table_obj is None or not hasattr(table_obj, 'matrix'):
            return
        raw_matrix = [[cell.value for cell in row] for row in table_obj.matrix]
        if len(raw_matrix) == 0:
            return

        rows = len(raw_matrix)
        cols = max(len(row) for row in raw_matrix)

        # Заголовки
        headers = []
        for cell_val in raw_matrix[0]:
            val_str = str(cell_val).strip() if pd.notna(cell_val) else ""
            if val_str.lower() in ["nan", "none"]:
                val_str = ""
            headers.append(val_str)

        if len(headers) < cols:
            headers.extend([""] * (cols - len(headers)))
        else:
            headers = headers[:cols]

        # Создаем таблицу в документе Word
        word_tab = self.doc.add_table(rows=1, cols=cols)
        word_tab.style = 'Table Grid'

        # Записываем шапку
        hdr_cells = word_tab.rows[0].cells
        for c in range(cols):
            hdr_cells[c].text = headers[c]
            for p in hdr_cells[c].paragraphs:
                for r in p.runs:
                    r.bold = True

        # Записываем тело данных
        for r in range(1, rows):
            row_cells = word_tab.add_row().cells
            for c in range(min(len(raw_matrix[r]), cols)):
                val = raw_matrix[r][c]
                if pd.isna(val) or str(val).lower() in ["nan", "none"]:
                    text_str = ""
                else:
                    text_str = str(round(val, 2)) if isinstance(val, (int, float)) else str(val)
                row_cells[c].text = text_str

    def AddPlot(self, figure, title="", width_inches=5.5):
        if figure is None:
            return

        buf = io.BytesIO()
        figure.savefig(buf, format='png', bbox_inches='tight', dpi=150)

        if title:
            self.doc.add_paragraph(title)

        self.doc.add_picture(buf, width=Inches(width_inches))
        plt.close(figure)

    def Save(self):
        n = 1
        while os.path.exists(f"{self.folder}/отчёт_{n}.docx"):
            n += 1
        filename = f"{self.folder}/отчёт_{n}.docx"

        self.doc.save(filename)
        return filename

    def BuildDocumentFromQueue(self, constructor_queue, plants_session_data, start_yr, end_yr, current_config):

        for elem in constructor_queue:
            plant_title = f"Завод: {elem['plant_name']} | " if elem['plant_name'] else ""
            self.AddHeading(f"{plant_title}{elem['label']}", level=1)
            current_plant_obj = None
            if elem['plant_name']:
                for p_id, p_info in plants_session_data.items():
                    if p_info["name"] == elem['plant_name']:
                        current_plant_obj = load_base_excel_files(p_info["folder_path"], p_info["name"])
                        if current_plant_obj:
                            current_plant_obj.CalculateMatrices(p_info["technology"], p_info["start_yr"],
                                                                p_info["end_yr"])
                        break

            # Сборка таблицы в отчет
            if elem["type"] == "table":
                target_table = None
                if elem["key"] == "global_t1" and current_plant_obj:
                    target_table = current_plant_obj.t1
                elif current_plant_obj:
                    if "series" in elem["field_name"]:
                        series_dict = getattr(current_plant_obj, elem["field_name"]).get(elem["tech"], {})
                        target_table = series_dict.get(elem["year"])
                    elif elem["field_name"] == "t7":
                        target_table = current_plant_obj.t7.get(elem["tech"])
                    else:
                        target_table = getattr(current_plant_obj, elem["field_name"])

                if target_table:
                    self.AddParagraph(getattr(target_table, 'description', ''))
                    self.AddTable(target_table)
                    self.AddPageBreak()
            # ВНУТРИ office.py (Сборка графиков по новой бэкенд-логике)
            elif elem["type"] == "graph" and current_plant_obj:
                years_range = list(range(start_yr, end_yr + 1))
                visualizer = NuclearDataVisualizer()

                if elem["field_name"] == "graph_t7":
                    # Забираем из logic.py две раздельные фигуры ПЗРО для вставки в Word
                    fig_w1, fig_w2 = visualizer.CreateBurialPlotsInCore(plants_session_data, start_yr, end_yr)
                    self.AddPlot(fig_w1, title="Динамика заполнения ПГЗР (Глубинное)", width_inches=elem["size"])
                    self.AddPlot(fig_w2, title="Динамика заполнения ППЗР (Приповерхностное)", width_inches=elem["size"])
                    self.AddPageBreak()


                # Отрисовка графика чувствительности (остается без изменений)
                elif elem["field_name"] == "graph_sens":
                    import streamlit as st
                    p_min_c = st.session_state.get('current_min_pct', 50)
                    p_max_c = st.session_state.get('current_max_pct', 150)
                    x_p, y_v = current_plant_obj.RunSensitivityAnalysis(current_config["technology"], start_yr, end_yr,
                                                                        p_min_c, p_max_c)
                    ax.plot(x_p, y_v, marker='s', color='#d9534f', linewidth=2)
                    ax.set_xlabel("Показатели")
                    ax.grid(True, linestyle='--', alpha=0.5)
                    self.AddPlot(fig, title=elem["label"], width_inches=elem["size"])
                    self.AddPageBreak()

    @staticmethod
    def RenderConstructorUI(start_yr, end_yr, current_config, active_id, technology):
        import streamlit as st
        import os
        from logic import ReportConstructor, NuclearDataVisualizer

        st.markdown("###Конструктор отчета Word (.docx)")
        st.caption("Выберите элементы программы, настройте их масштаб и порядок для записи в итоговый документ.")

        constructor = ReportConstructor()
        available_items = constructor.GetAvailableElements(st.session_state.plants_data, start_yr, end_yr)

        # Форма добавления элементов
        with st.container(border=True):
            st.markdown("**Шаг 1. Добавить элемент в документ**")
            item_options = {item["label"] + (f" ({item['group']})" if item['plant_name'] else ""): item for item in
                            available_items}
            selected_item_label = st.selectbox("Выберите таблицу или график для вставки:",
                                               options=list(item_options.keys()))
            element_size = st.slider("Ширина элемента в документе (дюймы):", min_value=2.0, max_value=8.0,
                                     value=5.5, step=0.5)

            if st.button("Добавить выбранный элемент в очередь отчета", use_container_width=True):
                chosen_meta = item_options[selected_item_label]
                st.session_state.constructor_queue.append({
                    "key": chosen_meta["key"], "type": chosen_meta["type"], "label": chosen_meta["label"],
                    "field_name": chosen_meta["field_name"], "plant_name": chosen_meta["plant_name"],
                    "tech": chosen_meta.get("tech"), "year": chosen_meta.get("year"), "size": element_size
                })
                st.toast("Элемент добавлен в очередь отчета!")
                st.rerun()

        # Отображение очереди документа
        st.markdown("**Шаг 2. Текущая структура документа (Очередь сборки)**")
        if not st.session_state.constructor_queue:
            st.info("Очередь пуста. Добавьте элементы с помощью формы выше.")
        else:
            for idx, q_item in enumerate(st.session_state.constructor_queue):
                q_col_text, q_col_size, q_col_del = st.columns([6.0, 2.0, 1.0])
                with q_col_text:
                    prefix = "таблица" if q_item["type"] == "table" else "график"
                    plant_prefix = f"[{q_item['plant_name']}] " if q_item["plant_name"] else "[Глобальный] "
                    st.write(f"{idx + 1}. {prefix} {plant_prefix}{q_item['label']}")
                with q_col_size:
                    st.caption(f"Ширина: {q_item['size']} дюйм.")
                with q_col_del:
                    if st.button("закрыть", key=f"del_q_{idx}"):
                        st.session_state.constructor_queue.pop(idx)
                        st.rerun()

            st.divider()

            # Кнопки действий
            btn_build_col, btn_cancel_col = st.columns(2)
            with btn_build_col:
                if st.button("Сгенерировать и сохранить отчет Word", use_container_width=True, type="primary"):
                    with st.spinner("Идет генерация документа конструктором..."):
                        from office import ReportDocument
                        doc_report = ReportDocument(title="Аналитический комплекс: Пользовательский отчет")
                        saved_file = doc_report.BuildDocumentFromQueue(
                            constructor_queue=st.session_state.constructor_queue,
                            plants_session_data=st.session_state.plants_data,
                            start_yr=start_yr, end_yr=end_yr, current_config=current_config
                        )
                        st.session_state.report_status = f"Пользовательский отчет '{os.path.basename(saved_file)}' успешно собран!"
                        st.session_state.constructor_queue = []
                        st.session_state.report_mode = False
                        st.rerun()
            with btn_cancel_col:
                if st.button("Закрыть конструктор (Вернуться к таблицам)", use_container_width=True):
                    st.session_state.report_mode = False
                    st.rerun()
